from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "superpowers" / "specs" / "2026-07-27-efficiency-report-generator-design.md"
OUTPUT = ROOT / "docs" / "效能报表自动生成组件需求规格说明书.docx"

FONT_CN = "Microsoft YaHei"
FONT_CODE = "Consolas"
NAVY = "17365D"
BLUE = "2F5597"
LIGHT_BLUE = "DCE6F1"
PALE_BLUE = "EEF3F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
DARK = "1F2937"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9460
TABLE_INDENT_DXA = 120


def set_run_font(run, name=FONT_CN, size=None, color=DARK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="C8D1DC", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MID_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=8.5, color=MID_GRAY)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 15, 7),
        ("Heading 2", 13, BLUE, 11, 5),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.46)
        style.paragraph_format.first_line_indent = Inches(-0.22)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.16

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = FONT_CODE
    code._element.rPr.rFonts.set(qn("w:ascii"), FONT_CODE)
    code._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CODE)
    code._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    code.font.size = Pt(8)
    code.font.color.rgb = RGBColor.from_string("243447")
    code.paragraph_format.left_indent = Inches(0.16)
    code.paragraph_format.right_indent = Inches(0.10)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.05
    p_pr = code._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)


def configure_section(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    left = p.add_run("效能报表自动生成组件")
    set_run_font(left, size=8.5, color=MID_GRAY, bold=True)
    p.add_run("\t")
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.2))
    right = p.add_run("需求规格说明书 · V1.0")
    set_run_font(right, size=8.5, color=MID_GRAY)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("需求规格说明书")
    set_run_font(r, size=12, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("效能报表自动生成组件")
    set_run_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("配置驱动 · 多 SQL 数据集 · Excel / Word 自动生成")
    set_run_font(r, size=13, color=MID_GRAY)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2100, 7360])
    set_table_borders(table, color="D5DCE5", size="4")
    set_repeat_table_header(table.rows[0])
    for col_idx, text in enumerate(("文档信息", "内容")):
        cell = table.cell(0, col_idx)
        set_cell_shading(cell, NAVY)
        p0 = cell.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(text)
        set_run_font(r0, size=9.5, color=WHITE, bold=True)
    rows = [
        ("文档版本", "V1.0"),
        ("文档状态", "需求确认稿"),
        ("编制日期", "2026-07-27"),
        ("适用技术栈", "Java 1.8 / Spring Boot 2.7 / MySQL 5.7"),
    ]
    for idx, (label, value) in enumerate(rows):
        row_idx = idx + 1
        set_cell_shading(table.cell(row_idx, 0), PALE_BLUE)
        p1 = table.cell(row_idx, 0).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        set_run_font(r1, size=9.5, color=NAVY, bold=True)
        p2 = table.cell(row_idx, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_run_font(r2, size=9.5, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("用途")
    set_run_font(r, size=10, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("作为设计、开发、测试和验收的共同依据。")
    set_run_font(r, size=10.5, color=DARK)
    doc.add_page_break()


def add_contents(doc, headings):
    p = doc.add_paragraph()
    r = p.add_run("目录")
    set_run_font(r, size=20, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(12)

    for heading in headings:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.2)
        p.paragraph_format.space_after = Pt(2.5)
        r = p.add_run(heading)
        set_run_font(r, size=10, color=DARK)
    doc.add_page_break()


def parse_inline(paragraph, text, size=None):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name=FONT_CODE, size=size or 9, color=NAVY)
            set_cell_shading_for_run(run, "EDF2F7")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size or 10.5, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size or 10.5)


def set_cell_shading_for_run(run, fill):
    r_pr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    r_pr.append(shd)


def add_markdown_table(doc, raw_lines):
    parsed = []
    for line in raw_lines:
        values = [cell.strip() for cell in line.strip().strip("|").split("|")]
        parsed.append(values)
    if len(parsed) >= 2 and all(re.fullmatch(r":?-{3,}:?", x) for x in parsed[1]):
        parsed.pop(1)
    if not parsed:
        return
    cols = max(len(row) for row in parsed)
    weights = []
    for idx in range(cols):
        lengths = [len(row[idx]) if idx < len(row) else 0 for row in parsed]
        weights.append(max(8, min(max(lengths), 36)))
    total = sum(weights)
    widths = [int(TABLE_WIDTH_DXA * w / total) for w in weights]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)

    table = doc.add_table(rows=len(parsed), cols=cols)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row_idx, values in enumerate(parsed):
        for col_idx in range(cols):
            cell = table.cell(row_idx, col_idx)
            if row_idx == 0:
                set_cell_shading(cell, NAVY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            text = values[col_idx] if col_idx < len(values) else ""
            r = p.add_run(text)
            set_run_font(
                r,
                size=8.5,
                color=WHITE if row_idx == 0 else DARK,
                bold=row_idx == 0,
            )
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def add_code_block(doc, lines):
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    for idx, line in enumerate(lines):
        r = p.add_run(line)
        set_run_font(r, name=FONT_CODE, size=8, color="243447")
        if idx < len(lines) - 1:
            r.add_break()


def build_document():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    body_start = next(i for i, line in enumerate(lines) if line.startswith("## 1."))
    body = lines[body_start:]
    headings = [line[3:].strip() for line in body if line.startswith("## ")]

    doc = Document()
    doc.core_properties.title = "效能报表自动生成组件需求规格说明书"
    doc.core_properties.subject = "配置驱动的效能报表 Excel/Word 自动生成组件"
    doc.core_properties.author = "项目需求组"
    doc.core_properties.keywords = "Java 1.8, Spring Boot 2.7, MySQL 5.7, Excel, Word, 报表"

    configure_styles(doc)
    configure_section(doc.sections[0])
    add_cover(doc)
    add_contents(doc, headings)

    i = 0
    in_code = False
    code_lines = []
    while i < len(body):
        line = body[i]
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, code_lines)
                in_code = False
                code_lines = []
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(body) and body[i].startswith("|"):
                table_lines.append(body[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
        elif re.match(r"^\d+\.\s+", line):
            text_value = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, text_value)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            p.paragraph_format.widow_control = True
            parse_inline(p, line.strip())
        i += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
