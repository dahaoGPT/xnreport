"""Build the committed Word template for the API efficiency example report."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "templates" / "api-design-efficiency.docx"
FONT_NAME = "宋体"
ACCENT = RGBColor(31, 78, 120)
MUTED = RGBColor(89, 89, 89)


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_fixed_table_width(table, widths_cm):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_dxa = int(sum(widths_cm) / 2.54 * 1440)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    widths_dxa = [int(round(width / 2.54 * 1440)) for width in widths_cm]
    widths_dxa[-1] = total_dxa - sum(widths_dxa[:-1])
    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            row.cells[index].width = Cm(width)
            set_cell_margins(row.cells[index])


def set_style_font(style, size, bold=False, color=None, before=0, after=6):
    style.font.name = FONT_NAME
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.1


def configure_styles(document):
    styles = document.styles
    set_style_font(styles["Normal"], 10.5, after=6)
    set_style_font(styles["Title"], 32, bold=True, color=RGBColor(0, 0, 0), after=20)
    heading_tokens = {
        1: (16, 16, 8),
        2: (14, 12, 6),
        3: (12, 8, 4),
        4: (11, 6, 3),
    }
    for level, (size, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        set_style_font(style, size, bold=True, color=ACCENT, before=before, after=after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.outline_level = level - 1
    if "TOC Heading" in styles:
        set_style_font(styles["TOC Heading"], 20, bold=True, color=RGBColor(0, 0, 0), after=18)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_toc(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-4" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在打开 Word 后自动更新"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def add_cover(document):
    for _ in range(3):
        document.add_paragraph()
    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(20)
    set_run_font(title.add_run("{{cover:title}}"), 32, True, RGBColor(0, 0, 0))

    organization = document.add_paragraph()
    organization.alignment = WD_ALIGN_PARAGRAPH.CENTER
    organization.paragraph_format.space_after = Pt(32)
    set_run_font(organization.add_run("——{{cover:organization}}"), 16, False, RGBColor(0, 0, 0))

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_fixed_table_width(table, [3.0, 9.0])
    labels = ["时　间", "{{cover:reportPeriod}}"]
    for index, value in enumerate(labels):
        cell = table.cell(0, index)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(value), 14)

    for _ in range(13):
        document.add_paragraph()
    prepared_by = document.add_paragraph()
    prepared_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared_by.paragraph_format.space_after = Pt(8)
    set_run_font(prepared_by.add_run("{{cover:preparedBy}}"), 13)
    prepared_date = document.add_paragraph()
    prepared_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(prepared_date.add_run("{{cover:preparedDate}}"), 13)
    prepared_date.add_run().add_break(WD_BREAK.PAGE)


def add_header_footer(document):
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.add_run("研发效能报告 | {{cover:organization}}"), 9, False, MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.add_run("第 "), 9, False, MUTED)
    add_page_number(paragraph)
    set_run_font(paragraph.add_run(" 页 | {{cover:preparedDate}}"), 9, False, MUTED)


def build():
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.start_type = WD_SECTION.NEW_PAGE
    configure_styles(document)
    add_header_footer(document)
    add_cover(document)

    toc_heading = document.add_paragraph()
    toc_heading.style = document.styles["TOC Heading"]
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading.paragraph_format.space_after = Pt(18)
    set_run_font(toc_heading.add_run("目　录"), 20, True, RGBColor(0, 0, 0))
    toc = document.add_paragraph()
    add_toc(toc)
    toc.add_run().add_break(WD_BREAK.PAGE)

    anchor = document.add_paragraph()
    anchor.paragraph_format.space_after = Pt(0)
    set_run_font(anchor.add_run("{{sections}}"), 10.5)

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    document.core_properties.title = "研发效能报告模板"
    document.core_properties.subject = "配置驱动研发效能报表"
    document.core_properties.author = "效能小组"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
